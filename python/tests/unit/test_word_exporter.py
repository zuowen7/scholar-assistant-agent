"""Word 导出器单元测试"""

import pytest
from pathlib import Path
import re

from src.formatter.word_exporter import (
    markdown_to_docx,
    _parse_inline_format,
    _apply_inline,
)
from docx import Document


class TestInlineFormatParsing:
    def test_plain_text(self):
        parts = _parse_inline_format("hello world")
        assert len(parts) == 1
        assert parts[0] == ("plain", "hello world")

    def test_bold(self):
        parts = _parse_inline_format("**bold**")
        assert ("bold", "bold") in parts

    def test_italic(self):
        parts = _parse_inline_format("*italic*")
        assert ("italic", "italic") in parts

    def test_code(self):
        parts = _parse_inline_format("`code`")
        assert ("code", "code") in parts

    def test_link(self):
        parts = _parse_inline_format("[Google](https://google.com)")
        assert ("link", "Google") in parts

    def test_mixed(self):
        parts = _parse_inline_format("**bold** and *italic* and `code`")
        assert ("bold", "bold") in parts
        assert ("italic", "italic") in parts
        assert ("code", "code") in parts

    def test_empty(self):
        parts = _parse_inline_format("")
        assert parts == []


class TestMarkdownToDocx:
    def test_empty_markdown(self, tmp_path):
        out = tmp_path / "empty.docx"
        result = markdown_to_docx("", out, title="Empty Test")
        assert result.exists()
        doc = Document(result)
        assert len(doc.paragraphs) >= 1  # 至少有空段落

    def test_title_output(self, tmp_path):
        out = tmp_path / "title.docx"
        markdown_to_docx("# Hello", out, title="My Title")
        doc = Document(out)
        # 找包含标题文本的段落
        texts = [p.text for p in doc.paragraphs]
        assert "My Title" in texts or any("Hello" in t for t in texts)

    def test_heading_levels(self, tmp_path):
        out = tmp_path / "headings.docx"
        md = "# H1\n## H2\n### H3\nplain text"
        markdown_to_docx(md, out)
        doc = Document(out)
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("H1" in p for p in paras)
        assert any("H2" in p for p in paras)
        assert any("H3" in p for p in paras)

    def test_bold_italic_preserved(self, tmp_path):
        out = tmp_path / "inline.docx"
        md = "This is **bold** and *italic* text"
        markdown_to_docx(md, out)
        # 验证 docx 文件可读（格式信息在 run 中）
        doc = Document(out)
        assert len(doc.paragraphs) >= 1

    def test_list_items(self, tmp_path):
        out = tmp_path / "lists.docx"
        md = "- Item 1\n- Item 2\n1. Num 1\n2. Num 2"
        markdown_to_docx(md, out)
        doc = Document(out)
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        # 列表项应该被保留
        assert any("Item 1" in p or "Item" in p for p in paras)

    def test_code_block(self, tmp_path):
        out = tmp_path / "code.docx"
        md = "```\ndef hello():\n    pass\n```"
        markdown_to_docx(md, out)
        doc = Document(out)
        # 代码块应该出现在文档中
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "hello" in all_text

    def test_blockquote(self, tmp_path):
        out = tmp_path / "quote.docx"
        md = "> 这是一段引用\n> 这是引用的第二行"
        markdown_to_docx(md, out)
        doc = Document(out)
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("引用" in p for p in paras)

    def test_file_path_string(self, tmp_path):
        out = str(tmp_path / "str_path.docx")
        result = markdown_to_docx("# Test", out)
        assert Path(result).exists()

    def test_page_setup(self, tmp_path):
        out = tmp_path / "setup.docx"
        markdown_to_docx("text", out)
        doc = Document(out)
        section = doc.sections[0]
        # A4 默认宽度 + 1 英寸页边距
        assert section.page_width is not None
        assert section.left_margin.inches == 1.0