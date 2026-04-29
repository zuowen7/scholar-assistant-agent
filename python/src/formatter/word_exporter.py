"""Word (.docx) 导出器 — 将 Markdown 文本转换为带格式的 .docx 文件

核心功能：
1. 解析 Markdown 标题层级 (# / ## / ###) → Word 内置 Heading 样式
2. 保留段落结构（双换行 = 段间距）
3. 处理加粗 (**text**)、斜体 (*text*)、行内代码 (`code`)
4. 识别链接、引用块、有序/无序列表、表格、水平线、代码块
5. A4 页面，学术论文排版（宋体正文 / 黑体标题）

依赖：python-docx >= 1.1.0
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 页面常量 ──────────────────────────────────────────────────────
A4_WIDTH = Cm(21.0)
A4_HEIGHT = Cm(29.7)
MARGIN = Cm(2.54)  # 1 inch

FONT_BODY = "宋体"
FONT_HEADING = "黑体"
FONT_CODE = "Consolas"
FONT_SIZE_BODY = Pt(12)     # 小四
FONT_SIZE_H1 = Pt(22)       # 二号
FONT_SIZE_H2 = Pt(16)       # 三号
FONT_SIZE_H3 = Pt(14)       # 四号
LINE_SPACING = 1.5


def _parse_inline_format(text: str) -> list[tuple[str, str]]:
    """解析行内格式 → [(type, content), ...]"""
    parts: list[tuple[str, str]] = []
    pattern = re.compile(
        r'`([^`]+)`|'                           # code
        r'\*\*([^*]+)\*\*|'                     # bold
        r'\*([^*]+)\*|'                         # italic
        r'___(.+?)___|'                         # bold+italic
        r'__([^_]+)__|'                         # bold (underscore)
        r'_([^_]+)_|'                           # italic (underscore)
        r'\[([^\]]+)\]\([^)]+\)|'               # link
        r'([^`*\[_]+)',                         # plain
        re.DOTALL
    )
    for m in pattern.finditer(text):
        if m.group(1):      parts.append(("code", m.group(1)))
        elif m.group(2):    parts.append(("bold", m.group(2)))
        elif m.group(3):    parts.append(("italic", m.group(3)))
        elif m.group(4):    parts.append(("bolditalic", m.group(4)))
        elif m.group(5):    parts.append(("bold", m.group(5)))
        elif m.group(6):    parts.append(("italic", m.group(6)))
        elif m.group(7):    parts.append(("link", m.group(7)))
        elif m.group(8):    parts.append(("plain", m.group(8)))
    return parts


def _apply_inline(paragraph, text: str) -> None:
    """在段落中应用行内格式"""
    for fmt, content in _parse_inline_format(text):
        run = paragraph.add_run(content)
        run.font.name = FONT_BODY
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), FONT_BODY)

        if fmt == "bold":
            run.bold = True
        elif fmt == "italic":
            run.italic = True
        elif fmt == "bolditalic":
            run.bold = True
            run.italic = True
        elif fmt == "code":
            run.font.name = FONT_CODE
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            hl = OxmlElement('w:highlight')
            hl.set(qn('w:val'), 'lightGray')
            rPr.append(hl)
        elif fmt == "link":
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)


def _set_run_font(run, name: str, size, color=None):
    """设置 run 的字体（含 eastAsia fallback）"""
    run.font.name = name
    run.font.size = size
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def _parse_table(table_text: str) -> list[list[str]]:
    """解析 markdown 表格文本 → 二维数组"""
    rows = []
    for line in table_text.strip().split('\n'):
        line = line.strip()
        if not line or re.match(r'^[\s|:-]+$', line):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    return rows


def _write_table(doc, table_text: str) -> None:
    """将 markdown 表格写入 Word"""
    rows = _parse_table(table_text)
    if len(rows) < 2:
        return
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < ncols:
                cell = table.cell(i, j)
                cell.text = ''
                p = cell.paragraphs[0]
                _apply_inline(p, cell_text)
                for run in p.runs:
                    run.font.size = Pt(10)
                if i == 0:
                    for run in p.runs:
                        run.bold = True
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'E8E8E8')
                    shading.set(qn('w:val'), 'clear')
                    cell._tc.get_or_add_tcPr().append(shading)


def _write_code_block(doc, lines: list[str]) -> None:
    """写入等宽代码块"""
    code_text = '\n'.join(lines)
    para = doc.add_paragraph()
    run = para.add_run(code_text)
    _set_run_font(run, FONT_CODE, Pt(9), RGBColor(0x2D, 0x2D, 0x2D))
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)


def _write_blockquote(doc, lines: list[str]) -> None:
    """写入引用块"""
    for line in lines:
        para = doc.add_paragraph()
        run = para.add_run(line)
        _set_run_font(run, FONT_BODY, Pt(12), RGBColor(0x55, 0x55, 0x55))
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '12')
        left.set(qn('w:space'), '8')
        left.set(qn('w:color'), '999999')
        pBdr.append(left)
        pPr.append(pBdr)
        para.paragraph_format.left_indent = Cm(1)
        para.paragraph_format.space_after = Pt(3)


def _write_list_item(doc, text: str, ordered: bool = False, level: int = 0) -> None:
    """写入列表项"""
    para = doc.add_paragraph()
    indent = Cm(1.27 * (level + 1))
    para.paragraph_format.left_indent = indent
    para.paragraph_format.first_line_indent = Cm(-0.63)
    marker = '•' if not ordered else ''
    _apply_inline(para, f'{marker} {text}' if marker else text)
    para.paragraph_format.space_after = Pt(2)


# ── 主函数 ──────────────────────────────────────────────────────

def markdown_to_docx(
    markdown_text: str,
    output_path: str | Path,
    *,
    title: str = "Scholar Assistant Export",
) -> Path:
    output_path = Path(output_path)
    # Strip XML-incompatible control characters (keep \n \r \t)
    markdown_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', markdown_text)
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = A4_WIDTH
    section.page_height = A4_HEIGHT
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN

    # ── Normal 样式 ──
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = FONT_SIZE_BODY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = LINE_SPACING
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_BODY)

    # ── 标题 ──
    if title:
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = tp.add_run(title)
        _set_run_font(tr, FONT_HEADING, FONT_SIZE_H1, RGBColor(0x00, 0x00, 0x00))
        tr.bold = True
        tp.paragraph_format.space_after = Pt(18)

    # ── 解析 ──
    lines = markdown_text.split('\n')
    i = 0
    in_code_block = False
    code_lines: list[str] = []
    in_table = False
    table_lines: list[str] = []
    ordered_counter = 0

    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                _write_code_block(doc, code_lines)
                in_code_block = False
                code_lines = []
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 表格行（以 | 开头且包含至少两个 |）
        if '|' in line and line.strip().startswith('|') and line.strip().endswith('|'):
            table_lines.append(line)
            in_table = True
            i += 1
            continue
        elif in_table:
            _write_table(doc, '\n'.join(table_lines))
            table_lines = []
            in_table = False
            # don't increment i — process current line normally

        # 标题
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            ordered_counter = 0
            level = min(len(m.group(1)), 3)
            text = m.group(2).strip()
            para = doc.add_paragraph()
            _apply_inline(para, text)
            for run in para.runs:
                run.bold = True
                _set_run_font(run, FONT_HEADING,
                              {1: FONT_SIZE_H1, 2: FONT_SIZE_H2, 3: FONT_SIZE_H3}.get(level, Pt(13)),
                              RGBColor(0x00, 0x00, 0x00))
            para.paragraph_format.space_before = Pt(18 if level == 1 else 12)
            para.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # 引用块
        if line.strip().startswith('>'):
            ordered_counter = 0
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                ql = lines[i].strip().lstrip('>').strip()
                if ql:
                    quote_lines.append(ql)
                i += 1
            _write_blockquote(doc, quote_lines)
            continue

        # 有序列表
        m = re.match(r'^\s*(\d+)\.\s+(.*)', line)
        if m:
            ordered_counter += 1
            text = f"{ordered_counter}. {m.group(2)}"
            _write_list_item(doc, text, ordered=True)
            i += 1
            continue

        # 无序列表
        m = re.match(r'^[-*+]\s+(.*)', line)
        if m:
            ordered_counter = 0
            _write_list_item(doc, m.group(1), ordered=False)
            i += 1
            continue

        # 水平线
        if re.match(r'^[-*_]{3,}$', line.strip()):
            ordered_counter = 0
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run('─' * 50)
            run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            run.font.size = Pt(8)
            i += 1
            continue

        # 普通段落
        if line.strip():
            ordered_counter = 0
            para = doc.add_paragraph()
            _apply_inline(para, line.strip())
            para.paragraph_format.space_after = Pt(6)
            pf = para.paragraph_format
            pf.line_spacing = LINE_SPACING
            # 首行缩进 2 字符
            pf.first_line_indent = Cm(0.74)
        i += 1

    # 收尾：还在表格中
    if in_table and table_lines:
        _write_table(doc, '\n'.join(table_lines))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
