"""Word (.docx) 导出器 — 将 Markdown 文本转换为带格式的 .docx 文件

核心功能：
1. 解析 Markdown 标题层级 (# / ## / ###) → Word 样式 (Heading 1/2/3)
2. 保留段落结构（双换行 = 段间空行，单换行 = 软换行）
3. 处理加粗 (**text**)、斜体 (*text*)、行内代码 (`code`)
4. 识别 Markdown 链接 [text](url) 并保留文本（超链接可选保留）
5. 处理引用块 (>) 和列表 (- / 1.)

依赖：python-docx >= 1.1.0
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Markdown 标记 → Word 样式映射
def _set_heading_style(paragraph: "paragraph", level: int) -> None:
    """设置标题样式，level 1=Heading1, 2=Heading2, 3=Heading3"""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    run.bold = True
    base_size = {
        1: 18,
        2: 16,
        3: 14,
    }.get(level, 14)
    run.font.size = Pt(base_size)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # 深蓝色


def _parse_inline_format(text: str) -> list[tuple[str, str]]:
    """解析行内格式，返回 (格式类型, 内容) 列表

    格式类型: "bold" | "italic" | "code" | "link" | "plain"
    """
    parts: list[tuple[str, str]] = []

    # 按优先级解析: 代码 > 粗体/斜体 > 链接 > 普通
    pattern = re.compile(
        r'`([^`]+)`|'                           # 行内代码
        r'\*\*([^*]+)\*\*|'                     # 粗体
        r'\*([^*]+)\*|'                         # 斜体
        r'\[([^\]]+)\]\([^)]+\)|'               # 链接 [text](url)
        r'([^`*\[]+)',                          # 普通文本
        re.DOTALL
    )

    for m in pattern.finditer(text):
        if m.group(1):      parts.append(("code", m.group(1)))
        elif m.group(2):    parts.append(("bold", m.group(2)))
        elif m.group(3):    parts.append(("italic", m.group(3)))
        elif m.group(4):    parts.append(("link", m.group(4)))
        elif m.group(5):    parts.append(("plain", m.group(5)))

    return parts


def _apply_inline_format(paragraph: "paragraph", text: str) -> None:
    """在段落中应用行内格式（加粗/斜体/代码/链接）"""
    parts = _parse_inline_format(text)
    for fmt_type, content in parts:
        run = paragraph.add_run(content)
        if fmt_type == "bold":
            run.bold = True
        elif fmt_type == "italic":
            run.italic = True
        elif fmt_type == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xD0, 0x33, 0x20)
            # 代码背景色（通过 highlight）
            rPr = run._r.get_or_add_rPr()
            highlight = OxmlElement('w:highlight')
            highlight.set(qn('w:val'), 'lightGray')
            rPr.append(highlight)
        elif fmt_type == "link":
            # 链接保持文本颜色，不设下划线（可选改为蓝色下划线）
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)


def markdown_to_docx(
    markdown_text: str,
    output_path: str | Path,
    *,
    title: str = "研墨导出",
    page_width: float = 6.5,  # inch, 正文宽度（留边距）
    font_name: str = "宋体",   # 正文字体（中文）
    code_font: str = "Consolas",
    font_size: float = 11.0,  # 正文字号（pt）
    heading_font: str = "黑体",
) -> Path:
    """将 Markdown 文本转换为带格式的 .docx 文件

    Args:
        markdown_text: Markdown 格式文本
        output_path: 输出 .docx 路径
        title: 文档标题（写入会作为首行大标题）
        page_width: 页面宽度（英寸），默认 6.5 留出 1 英寸边距
        font_name: 中文正文字体
        code_font: 代码字体
        font_size: 正文字号
        heading_font: 标题字体

    Returns:
        输出的 Path 对象
    """
    output_path = Path(output_path)
    doc = Document()

    # ── 页面设置 ──────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.5)          # A4 宽度
    section.page_height = Inches(11.69)       # A4 高度
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ── 默认样式 ──────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # 如果 Normal 样式没有中文字体兼容，发一篇默认段落
    doc.add_paragraph()

    # ── 标题 ──────────────────────────────────────────────────────
    if title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.name = heading_font
        title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        title_para.paragraph_format.space_after = Pt(18)

    # ── 处理段落 ──────────────────────────────────────────────────
    lines = markdown_text.split("\n")
    i = 0
    in_code_block = False
    code_block_lines: list[str] = []
    in_blockquote = False

    while i < len(lines):
        line = lines[i]

        # 代码块: ```...```
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_block_lines = []
            else:
                # 代码块结束，写入
                _write_code_block(doc, code_block_lines, code_font)
                in_code_block = False
                code_block_lines = []
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Markdown 标题: # ## ###
        m = re.match(r'^(#{1,3})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()

            # Word 中没有内置 Heading 样式时，手动创建
            h_para = doc.add_paragraph()
            h_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            h_run = h_para.add_run(text)
            h_run.bold = True
            h_run.font.name = heading_font
            h_run.font.size = Pt({1: 18, 2: 16, 3: 14}.get(level, 14))
            h_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            h_para.paragraph_format.space_before = Pt(12)
            h_para.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # 引用块: > text
        if line.strip().startswith(">"):
            # 收集连续引用行
            quote_lines = []
            while i < len(lines) and (lines[i].strip().startswith(">") or not lines[i].strip()):
                ql = lines[i].strip().lstrip(">").strip()
                if ql:
                    quote_lines.append(ql)
                i += 1
                # 遇到非引用行且非空行则停止
                if i < len(lines) and not lines[i].strip().startswith(">") and lines[i].strip():
                    break

            _write_blockquote(doc, quote_lines, font_name, font_size)
            continue

        # 有序列表: 1. xxx
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            _write_list_item(doc, f"{m.group(1)}. {m.group(2)}", font_name, font_size, ordered=True)
            i += 1
            continue

        # 无序列表: - xxx 或 * xxx
        m = re.match(r'^[-*]\s+(.*)', line)
        if m:
            _write_list_item(doc, m.group(1), font_name, font_size, ordered=False)
            i += 1
            continue

        # 水平线: --- 或 *** 或 ___
        if re.match(r'^[-*_]{3,}$', line.strip()):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # 普通段落
        if line.strip():
            para = doc.add_paragraph()
            _apply_inline_format(para, line.strip())
            para.paragraph_format.space_after = Pt(6)
        else:
            # 空行：段间空行（Word 默认 paragraph spacing 控制）
            pass
        i += 1

    # ── 保存 ──────────────────────────────────────────────────────
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _write_code_block(doc: "Document", lines: list[str], font_name: str) -> None:
    """写入等宽代码块"""
    code_text = "\n".join(lines)
    para = doc.add_paragraph()
    run = para.add_run(code_text)
    run.font.name = font_name
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xD0, 0x33, 0x20)
    # 浅灰色背景
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)


def _write_blockquote(doc: "Document", lines: list[str], font_name: str, font_size: float) -> None:
    """写入引用块"""
    for line in lines:
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
        # 左边框
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '12')
        left.set(qn('w:space'), '8')
        left.set(qn('w:color'), '7F7F7F')
        pBdr.append(left)
        pPr.append(pBdr)
        para.paragraph_format.left_indent = Inches(0.2)
        para.paragraph_format.space_after = Pt(3)


def _write_list_item(
    doc: "Document",
    text: str,
    font_name: str,
    font_size: float,
    ordered: bool = False,
) -> None:
    """写入列表项"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.first_line_indent = Inches(-0.3)
    _apply_inline_format(para, text)
    para.paragraph_format.space_after = Pt(2)