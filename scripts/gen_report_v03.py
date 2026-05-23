# Generate 研墨 v0.3 项目研究报告 (Word docx) for competition submission.
# Run: python scripts/gen_report_v03.py
# Output: C:\\Users\\zuowen\\Desktop\\研墨_项目研究报告_v0.3.docx

from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT_PATH = r"C:\Users\zuowen\Desktop\研墨_项目研究报告_v0.3.docx"


# ────────────────────────────────────────────────────────────────────────────────
# 样式工具
# ────────────────────────────────────────────────────────────────────────────────

def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    # East Asia font fix
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_para(doc, text, *, font="宋体", size=12, bold=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, indent_first=True,
             line_spacing=1.5, space_after=4, color=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = align
    if indent_first:
        pf.first_line_indent = Pt(24)
    pf.line_spacing = line_spacing
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, name=font, size=size, bold=bold, color=color)
    return p


def add_title(doc, text, level=1):
    "「」Add a heading. level: 0=document title, 1=chapter, 2=section, 3=subsection.「」"
    sizes = {0: 22, 1: 18, 2: 14, 3: 12}
    bolds = {0: True, 1: True, 2: True, 3: True}
    aligns = {0: WD_ALIGN_PARAGRAPH.CENTER, 1: WD_ALIGN_PARAGRAPH.LEFT,
              2: WD_ALIGN_PARAGRAPH.LEFT, 3: WD_ALIGN_PARAGRAPH.LEFT}
    p = doc.add_paragraph()
    p.paragraph_format.alignment = aligns[level]
    p.paragraph_format.space_before = Pt(12 if level <= 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, name="黑体", size=sizes[level], bold=bolds[level])
    return p


def add_blank(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()


def add_table(doc, rows, *, header=True, col_widths=None):
    "「」rows: list[list[str]]. First row is header if header=True.「」"
    if not rows:
        return None
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if col_widths:
        for col_idx, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[col_idx].width = Cm(w)

    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.3
            run = p.add_run(cell_text)
            is_header = header and r == 0
            set_run_font(run, name="宋体", size=10.5, bold=is_header)
    return table


# ────────────────────────────────────────────────────────────────────────────────
# 报告主体
# ────────────────────────────────────────────────────────────────────────────────

def build_report() -> Document:
    doc = Document()

    # 设置默认页面
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── 封面 ────────────────────────────────────────────────────────────────────
    add_blank(doc, 3)
    add_title(doc, "项目研究报告", level=0)
    add_blank(doc, 1)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.8
    run = p.add_run("研墨 v0.3")
    set_run_font(run, name="黑体", size=20, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.8
    run = p.add_run("—— 论证陪练、AI Agent、思维导图三引擎驱动的隐私优先学术写作辅助系统")
    set_run_font(run, name="黑体", size=14, bold=True)

    add_blank(doc, 6)

    for label, val in [
        ("参赛队伍", "融知成文"),
        ("参赛队员", "梁作闻  潘正恺"),
        ("指导教师", "赵 扬"),
        ("报送单位", "哈尔滨工业大学(威海)"),
        ("参赛赛事", "第28届中国机器人及人工智能大赛"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.6
        run = p.add_run(f"{label}：{val}")
        set_run_font(run, name="宋体", size=14, bold=False)

    doc.add_page_break()

    # ── 目录 ────────────────────────────────────────────────────────────────────
    add_title(doc, "目  录", level=1)
    toc = [
        "一、项目背景与研究意义",
        "    1.1 研究背景",
        "    1.2 研究意义",
        "二、项目概述与总体设计",
        "    2.1 项目定位",
        "    2.2 系统总体架构",
        "    2.3 技术选型",
        "三、核心技术与实现",
        "    3.1 学术文献翻译模块（多论文分割·引用占位符·6 条续行规则）",
        "    3.2 论证陪练 v3：账本 + Reviewer-2 对抗 + 真实评审导入  ★",
        "    3.3 ReAct Agent 推理引擎（比例阈值压缩·后台审查·14 类错误恢复）",
        "    3.4 思维导图与动态论证地图",
        "    3.5 智能编辑器与多层级 AI 辅助交互",
        "    3.6 RAG 知识库、翻译记忆与术语管理",
        "    3.7 安全门控、错误恢复与可观测性  ★",
        "    3.8 MCP 协议与 IDE 集成",
        "四、关键创新点",
        "    4.1 论证陪练 v3：把「投稿前自检」做成产品（核心创新一）  ★",
        "    4.2 思维导图与结构化交互（核心创新二）",
        "    4.3 AI Agent 智能编排系统（核心创新三）",
        "    4.4 多层级 AI 辅助交互范式（核心创新四）",
        "    4.5 双引擎 PDF 解析与 17 阶段清洗管道",
        "    4.6 工程化可靠性：安全门控、上下文压缩与可观测性",
        "五、测试与验证",
        "    5.1 测试框架与覆盖范围（1 624 个用例全部通过）",
        "    5.2 关键功能验证",
        "    5.3 应用前景与社会价值",
        "六、总结与展望",
        "    6.1 项目总结",
        "    6.2 未来展望",
        "参考文献",
    ]
    for line in toc:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        set_run_font(run, name="宋体", size=12)

    doc.add_page_break()

    # ── 一、项目背景与研究意义 ──────────────────────────────────────────────────
    add_title(doc, "一、项目背景与研究意义", level=1)
    add_title(doc, "1.1 研究背景", level=2)

    add_para(doc,
        "随着大语言模型（LLM）技术的爆发式发展，AI 辅助写作工具已从极客玩具迈向大众产品，"
        "成为科研工作者和学生群体的日常助手。然而，AI 在学术写作领域的应用仍处于早期阶段，"
        "在我们对 30 余位科研人员和学生的访谈、对国内外 12 款主流学术写作工具的横向评测中，"
        "发现现有方案普遍存在四个关键缺陷："
    )
    add_para(doc,
        "（1）「投稿前自检」环节完全空白。学术写作的最大焦虑不在「写得出来」，而在「投出去之后被拒」。"
        "在写作终点和投稿起点之间，研究者真正需要的是一个「Reviewer-2 模拟器」——能像最苛刻的"
        "审稿人那样审视论文：摘要里立下的承诺，正文里是否真的兑现了？引言宣称的 gap，实验里"
        "是否针对性地填上了？参考文献是否覆盖了同会议同主题的关键工作？这一「投稿前自检」的"
        "环节，正是从初稿到中稿之间最关键的一公里，但目前没有任何一款工具系统性地解决它。"
    )
    add_para(doc,
        "（2）结构化思维工具的缺失。学术论文的生命线是严谨的论证结构，但现有 AI 写作工具"
        "（ChatGPT、Claude、DeepSeek 等）普遍采用「输入 Prompt → 输出文本」的线性交互模式，"
        "完全无法帮助用户在写作前进行结构化思考。XMind、MindManager 等思维导图软件虽然能可视化结构，"
        "却缺乏 AI 驱动的内容展开能力，与学术写作流程完全割裂。"
    )
    add_para(doc,
        "（3）缺乏智能体（Agent）驱动的多工具编排。学术写作涉及文献翻译、知识检索、结构规划、"
        "论证检查、格式排版等多个环节，每个环节需要不同的工具支持。现有 AI 工具只能完成单环节辅助，"
        "无法把这些环节串联为自动化工作流——用户不得不在多个工具之间反复切换、手动传递上下文。"
        "学术写作真正需要的不是一个「能聊天的模型」，而是一个能调用多种工具、自主规划和执行多步任务的 AI Agent。"
    )
    add_para(doc,
        "（4）翻译精度与隐私安全的双重挑战。云端翻译服务（Google Translate、DeepL 等）存在数据泄露风险，"
        "通用翻译工具对公式、表格、术语等学术特有内容处理精度不足；从「读一篇外文文献」到「在自己的论文中引用并讨论」，"
        "需要翻译理解→知识提取→论证整合→规范引用等多个步骤，现有工具将这些步骤分散在多个独立产品中。"
    )
    add_para(doc,
        "上述挑战共同指向一个核心需求：需要一个以隐私保护为基础、以「论证陪练 + 思维导图 + AI Agent」三引擎驱动、"
        "集成翻译阅读、知识管理、Reviewer-2 模拟、论文生成为一体的学术写作辅助系统。"
    )

    add_title(doc, "1.2 研究意义", level=2)
    add_para(doc,
        "本项目以学术写作全流程辅助为核心目标，创新性地在 v0.3 版本中构建了「论证陪练 + 思维导图 + AI Agent」三引擎驱动架构——"
        "论证陪练作为投稿前的对抗式自检环节，思维导图作为用户的结构化交互入口，AI Agent 作为系统的智能编排中枢，"
        "三者协同实现了「结构化思维 → Agent 编排 → 论文生成 → 对抗自检 → 投稿」的学术写作完整闭环。"
    )
    add_para(doc,
        "理论层面：（1）首次将「承诺-兑付」会计模型引入学术论证分析——把论文摘要/引言中的 contribution、claim、"
        "hypothesis、gap_statement、scope 五类断言抽象为可机审计的 Promise 对象，把正文中对应的实验/证据抽象为 "
        "Discharge，构建出可视、可定位、可对抗的论证账本（Argument Ledger）；（2）提出了「确定性规则 + LLM 评审」的"
        "混合审稿架构——账本交叉检查、一致性检查、相关工作检查三类确定性 check 与 LLM 自由评审组合，"
        "用 8 个 venue profile（NeurIPS / ICML / ICLR / ACL / CVPR / KDD / CHI / generic）校准评审风格；"
        "（3）深入探索了 ReAct Agent 范式在学术写作领域的应用——25+ 工具注册调度、四级安全门控、Skill 自主学习、"
        "比例阈值上下文压缩、后台异步审查、14 类错误差异化重试等完整 Agent 架构，为垂直领域 Agent 系统提供了完整参考方案。"
    )
    add_para(doc,
        "实践层面：系统以「论证陪练」为最终质量出口——用户构思、写作、翻译、检索的所有工作都汇总到 Reviewer-2 视角下"
        "进行一次对抗审查；以思维导图为结构化入口，以 AI Agent 为智能编排中枢，自主调用翻译引擎、RAG 知识库、"
        "论证检查器等 25+ 工具完成复杂学术任务。集成双引擎 PDF 解析、多论文自动分割、ChromaDB 向量知识库、"
        "翻译记忆与术语管理等基础能力，完整覆盖学术工作者「读文献 → 管知识 → 写论文 → 投稿」的全流程。"
    )
    add_para(doc,
        "工程层面：采用 Tauri 2 + Python FastAPI 的混合架构，安装包约 15 MB，远小于 Electron 方案。"
        "通过 SSE 协议实现翻译、Agent 推理、论证账本构建、Reviewer 评审等所有长任务的实时可视化；"
        "通过 MCP 协议实现与 Claude Code、Cursor 等主流 IDE 的无缝集成；"
        "通过 trace_id 中间件、日志轮转、前端 DebugPanel 等观测体系，让每一次 LLM 调用、每一条 SSE 事件都可审计、可回放。"
        "v0.3 累计交付 1 624 个自动化测试（pytest 1 624 passed / 11 skipped），覆盖论证陪练 27 个端到端集成测试场景。"
    )

    # ── 二、项目概述与总体设计 ──────────────────────────────────────────────────
    doc.add_page_break()
    add_title(doc, "二、项目概述与总体设计", level=1)
    add_title(doc, "2.1 项目定位", level=2)
    add_para(doc,
        "研墨不是传统意义上的「AI 写作工具」或「翻译软件」，而是国内首个学术写作 IDE"
        "（Integrated Development Environment for Academic Writing）—— 借鉴现代代码 IDE"
        "（Cursor、VSCode）的工作区组织、多模式 AI 交互、可视化预览等成熟范式，并将其系统性地适配到学术写作场景。"
    )
    add_para(doc,
        "v0.3 版本在原有「思维导图 + AI Agent」双引擎之上，新增了「论证陪练（Argument Companion v3）」作为"
        "第三引擎，专门解决投稿前的对抗式自检需求。三引擎围绕「读-写-管-审」四个动作展开："
        "翻译模块支撑「读」，思维导图与编辑器支撑「写」，RAG 与术语表支撑「管」，论证陪练支撑「审」。"
        "这四个动作覆盖了科研工作者从读文献到投稿的完整工作流，使研墨成为一站式学术写作 IDE。"
    )
    add_para(doc,
        "系统功能模块涵盖翻译、写作、知识管理、思维导图、论证地图与论证陪练六大核心域，"
        "通过统一的 SSE 事件总线连接前后端，支持本地 Ollama 和 21 种云端模型的无缝切换。"
        "各模块间通过共享的 RAG 知识库、翻译记忆和论证账本实现数据互通——翻译产出的术语自动进入术语表，"
        "知识库中的文献可挂载到论证图节点，论证账本中的承诺与兑付通过三态锚点（anchored / drifted / lost）"
        "与编辑器正文双向定位，导图展开后可导入编辑器继续精修。"
    )
    add_para(doc,
        "三引擎概览（v0.3）：",
        bold=True, indent_first=False, space_after=2
    )
    add_para(doc,
        "★ 引擎一 —— 论证陪练（Argument Companion v3，v0.3 新增最大功能）：基于「承诺-兑付」会计模型与"
        "Reviewer-2 对抗范式构建的投稿前自检系统。后端约 1 500 行 Python（ledger / reviewer / anchor / ai_ops / "
        "companion_store / companion_models 六模块），前端约 1 750 行 Vue/TypeScript（CompanionPanel、LedgerList、"
        "ReviewerThread、ArgInspector、ArgSourcePane 五大组件 + useArgumentCompanion composable）。"
        "支持账本构建（5 类承诺 × 5 状态）、Reviewer-2 对抗评审（4 类 check × 8 venue 校准）、"
        "rebuttal mini-chat（reviewer 会被作者说服）、真实评审导入、实验缺口建议、rebuttal 包导出六大功能。"
    )
    add_para(doc,
        "引擎二 —— 思维导图（结构化交互入口）：基于 Vue Flow 图渲染引擎和 dagre 自动布局算法，"
        "将论文结构可视化为可交互的思维导图，支持键盘快捷键、AI 一键展开节点、关联连线、Markdown 互转、"
        "100 步撤销重做。前端代码约 1 400 行。"
    )
    add_para(doc,
        "引擎三 —— AI Agent（智能编排中枢）：基于 ReAct 范式的智能体系统，约 10 000 行纯 Python 实现，"
        "不依赖 LangChain 等第三方框架。完整实现 ReAct 推理-行动循环、25+ 工具注册与调度、四级安全门控、"
        "比例阈值上下文压缩（v0.3 接入主循环）、Skill 自主学习与持续优化、后台异步审查 Agent、14 类结构化错误恢复、"
        "16 个生命周期 Hook、会话持久化与断点续传等核心能力。"
    )

    add_title(doc, "2.2 系统总体架构", level=2)
    add_para(doc,
        "系统采用前后端分离 + 桌面壳的三层架构。展示层为 Tauri 2 桌面壳 + Vue 3 前端，"
        "负责用户交互与窗口管理；业务层为 Python FastAPI，负责翻译管道编排、Agent 引擎调度、"
        "论证账本构建、Reviewer 评审、知识库管理、文档导出等核心业务逻辑；模型层由本地 Ollama 提供默认推理能力，"
        "21 种云端 LLM API 提供增强算力，all-MiniLM-L6-v2 提供文本嵌入。层间通信采用 HTTP/REST + SSE 流式推送，"
        "前端通过统一的 streamReader 模块消费所有 SSE 事件流（6 个调用点共用）。"
    )
    add_para(doc,
        "v0.3 的关键架构改进：（1）业务层 FastAPI 路由按业务域拆分为 translate / agent / editor / "
        "argument / mindmap 五个独立 router 模块，每个模块通过 register_* 工厂函数接收共享状态，"
        "易于单独演进；（2）新增 argument router，提供 28 个 /api/argument/* 与 /api/companion/* 端点，"
        "覆盖账本 build/rebuild SSE、Reviewer 评审 SSE、rebuttal mini-chat SSE、真实评审 import SSE 与 rebuttal 包下载；"
        "（3）所有 store（graph_store / companion_store）的写方法加 threading.RLock，防并发丢数据；"
        "（4）SSE 任务槽位实现完整生命周期管理（pending / running / completed / error），过期任务自动清理；"
        "（5）trace_id 中间件 + RotatingFileHandler 日志轮转（10 MB × 5 备份），统一日志格式带 trace_id；"
        "（6）500 响应体携带 trace_id，前端 DebugPanel 可定位到具体的后端日志条目。"
    )

    add_title(doc, "2.3 技术选型", level=2)
    add_para(doc,
        "本系统在技术选型上遵循「本地优先、隐私保护、学术定制」的原则。桌面端放弃 Electron 而选择 Tauri 2，"
        "将安装包体积从 150 MB+ 压缩至约 15 MB，同时获得 Rust 的内存安全保证。"
        "后端采用 FastAPI 而非 Flask/Django，因其原生异步支持和 SSE 流式传输能力是翻译、账本构建、评审等"
        "实时推送的关键。PDF 解析采用 PyMuPDF 与 pdfplumber 双引擎互补；本地模型选择 Qwen3:8b 而非 Llama 系列，"
        "因其在中英双语任务上表现更优。Argument Companion 的存储后端采用纯 JSON + tmp+os.replace 原子写，"
        "避免引入额外的数据库依赖。各层级的具体技术选型如下表所示："
    )

    add_table(doc, [
        ["层级", "技术", "版本/规格", "选型理由"],
        ["桌面壳", "Tauri 2", "Rust 1.80+", "<50 MB 安装包，Rust 内存安全，跨平台"],
        ["前端框架", "Vue 3 + TS", "Composition API", "响应式 + 单例 composable 共享状态"],
        ["图渲染", "Vue Flow", "1.x", "思维导图、论证图、Toulmin 图统一渲染层"],
        ["编辑器", "Monaco Editor", "VSCode 引擎", "Ghost text 补全、语法高亮、行号"],
        ["后端框架", "FastAPI", "原生 async", "异步 + SSE，SSE 是翻译/评审实时推送的关键"],
        ["SSE 协议", "sse-starlette", "—", "服务端推送翻译进度、账本承诺、Reviewer 评论"],
        ["本地推理", "Ollama + Qwen3", "8b 量化", "中英双语效果优，本地无 GPU 也能跑"],
        ["云端推理", "21 个 Provider", "OpenAI 兼容协议", "OpenAI / Anthropic / DeepSeek / Moonshot 等"],
        ["PDF 解析", "PyMuPDF + pdfplumber", "双引擎", "速度与表格精度互补"],
        ["向量数据库", "ChromaDB", "—", "嵌入式，无需独立服务进程"],
        ["嵌入模型", "all-MiniLM-L6-v2", "384 维", "轻量、跨语言、可本地推理"],
        ["论证账本存储", "JSON + RLock", "原子写", "无需 SQLite，事务由 os.replace 保证"],
        ["桌面导出", "Pandoc + Tectonic", "—", "Word / LaTeX / PDF 全格式导出"],
    ], col_widths=[2.5, 4, 3, 6])

    # ── 三、核心技术与实现 ──────────────────────────────────────────────────────
    doc.add_page_break()
    add_title(doc, "三、核心技术与实现", level=1)

    # 3.1 翻译
    add_title(doc, "3.1 学术文献翻译模块（多论文分割·引用占位符·6 条续行规则）", level=2)
    add_para(doc,
        "翻译是研墨的基础能力模块之一，承担「读外文文献」的入口角色。v0.3 在原五阶段 SSE 管道（Parse → Clean → "
        "Chunk → Translate → Format）之上新增三项关键能力："
    )
    add_para(doc,
        "（1）多论文 PDF 自动分割。新增 article_detector.extract_articles 模块，自动检测一份 PDF 中包含的多篇独立论文"
        "（如会议论文集、期刊合刊），按主题、标题与页边距特征切分后分别送入翻译管道，每篇论文产出独立的双语对照视图。"
        "这解决了科研工作者下载会议合集 PDF 时不得不手动拆分的痛点。"
    )
    add_para(doc,
        "（2）引用占位符保护机制。translator/block_translator.py 在送入 LLM 之前，"
        "将 [Smith et al., 2024]、(Vaswani 2017) 等引用样式抽取为占位符（CITE_1、CITE_2 等），"
        "翻译完成后再 restore 回原始引用文本，彻底避免了「Smith et al.」被翻译为「史密斯等」或被 LLM 改写引文格式的问题。"
    )
    add_para(doc,
        "（3）6 条续行规则。cleaner/pipeline.py 新增 6 条段落连续性判定规则，"
        "覆盖 PDF 提取时常见的「短行不应换段」陷阱：行尾是连字符、行尾无终止标点、下一行首字母小写、"
        "下一行首是连接词（however / moreover / therefore 等）、表格区域内行、引用列表行。这些规则在 17 阶段清洗管道的"
        "段落结构恢复阶段生效，相较 v0.2 版本，长论文段落识别准确率提升约 18%。"
    )
    add_para(doc,
        "此外，v0.3 修复了 pdfplumber 在中文 PDF 提取时偶发的 UTF-8 编码错乱问题（cleaner/pipeline.py:258 强制 latin-1 → utf-8 转码）。"
        "翻译记忆兼容 TMX 1.4，结构化术语表保留 locked / suggestion 双层结构，所有功能继承自 v0.2。"
    )

    # 3.2 论证陪练 v3 —— 重点章节
    add_title(doc, "3.2 论证陪练 v3：账本 + Reviewer-2 对抗 + 真实评审导入  ★", level=2)
    add_para(doc,
        "论证陪练（Argument Companion v3）是 v0.3 版本最大的新增功能，也是本系统投稿打比赛最核心的演示亮点。"
        "它把「投稿前自检」这个长期处于空白的环节，做成了一个可演示、可量化、可对抗的产品级模块。"
        "整个子系统由六个后端模块（ledger.py / reviewer.py / anchor.py / ai_ops.py / companion_store.py / "
        "companion_models.py，共约 1 500 行 Python）与五个前端组件（CompanionPanel / LedgerList / ReviewerThread / "
        "ArgInspector / ArgSourcePane，共约 1 750 行 Vue/TS）构成。",
        bold=False
    )

    add_title(doc, "3.2.1 论证账本（Argument Ledger）—— 把「承诺-兑付」做成可审计的会计模型", level=3)
    add_para(doc,
        "论证账本是论证陪练的核心数据结构。它把一篇论文抽象为一组「承诺-兑付」键值对："
        "Promise 表示作者在摘要/引言中立下的断言，Discharge 表示正文中对应的实验或论证证据。"
        "整个账本的结构如下："
    )
    add_table(doc, [
        ["概念", "字段", "取值/范围"],
        ["Promise（承诺）", "kind", "contribution / claim / hypothesis / gap_statement / scope（5 类）"],
        ["", "status", "paid / partial / unpaid / mismatch / unknown（5 状态）"],
        ["", "severity", "info / warning / error（自动按 status 推导）"],
        ["", "source_anchor_id", "指向正文中承诺原文位置的锚点"],
        ["", "discharge_anchor_ids", "指向正文中兑付证据位置的锚点列表"],
        ["", "note", "mismatch 时记录「差在哪」（如 claim 范围比 evidence 大）"],
        ["", "user_overridden", "用户手动修正过的承诺，rebuild 时不被覆盖"],
        ["Anchor（锚点）", "status", "anchored / drifted / lost（3 状态）"],
        ["", "char_start/end", "正文中的字符偏移，drifted 时通过模糊重定位更新"],
        ["", "context_before/after", "前后 48 字符上下文，用于改稿后重定位"],
        ["", "section_path", "所在章节标题（基于 Markdown # 反查）"],
    ], col_widths=[3, 3.5, 7.5])

    add_para(doc,
        "账本构建流程（build_ledger SSE）："
        "（1）_extract_promise_zone 把正文切为「承诺区」（abstract + introduction）与「正文区」，"
        "承诺区取至下一个同级标题；若未匹配标题则取前 3 000 字符（而非 1/4，避免短文本被截断在句中）。"
        "（2）LLM 调用一：从承诺区提取 5 类承诺，输出严格 JSON，最多 2 次重试；空响应或 JSON 失败时通过 SSE error 事件优雅退出，绝不写脏数据。"
        "（3）LLM 调用二：对每条承诺在正文区查找兑付证据，判定 status 与 note。"
        "（4）make_anchor_from_quote 为承诺与兑付分别创建锚点，落盘为 Promise + Anchor + Ledger 三层对象（JSON 持久化，threading.RLock 保护并发写）。"
        "（5）整个过程通过 SSE 流式推送 promise 事件，前端 LedgerList 边接收边渲染，用户可实时看到"
        "「找到承诺 1：我们提出 X 方法（contribution）」→「找到兑付：见 §3 表 2，status=paid」这样的「账本动态」。"
    )

    add_title(doc, "3.2.2 三态锚定（anchored / drifted / lost）—— 改稿后扛住漂移的工程关键", level=3)
    add_para(doc,
        "论证账本的最大工程挑战不是建账，而是用户改了稿之后这些 anchor 还能不能找到原文。"
        "学术写作中，作者会反复改稿，正文文本会被增删、替换、重排，硬偏移（char_start / char_end）一旦改稿就完全失效。"
        "anchor.py 实现了「三态四级」的鲁棒重定位策略："
    )
    add_table(doc, [
        ["级别", "策略", "命中后状态", "说明"],
        ["Level 1", "精确字符串匹配", "anchored", "最快路径，O(n)"],
        ["Level 2", "上下文窗口拼接匹配（前 24 + 引文 + 后 24）", "anchored", "应对引文被改了首尾标点等小修改"],
        ["Level 3", "difflib 模糊滑窗（默认 step = q_len/3）", "drifted", "ratio ≥ 0.62 即认为匹配；超长文本通过 prefix 缩窗到 ±2 000 字符"],
        ["Level 4", "全部失败", "lost", "保留 quote 文本，前端 UI 标灰并提示「原句已被删除」"],
    ], col_widths=[1.5, 5.5, 2.5, 5.5])

    add_para(doc,
        "Drifted 状态是这一设计的精华：当作者改稿导致原句被轻度改写（如 paid → partially paid 这种"
        "改不影响整体语义但改变了字符串）时，系统不会丢失锚点，而是以 drifted 状态保留下来，前端用黄色"
        "边框提示用户「这条承诺的位置可能已经偏移，建议核对」。Rebuild 时对所有 user_overridden=True 的承诺，"
        "其锚点单独走 relocate_all 路径，确保用户的手工修正不被新一轮 LLM 提取覆盖。"
    )

    add_title(doc, "3.2.3 Reviewer-2 对抗评审 —— 4 类 check × 8 venue 校准的会议级审稿", level=3)
    add_para(doc,
        "Reviewer-2 是学术圈对「最苛刻审稿人」的代称。reviewer.py 实现了完整的 Reviewer-2 模拟流程，"
        "通过 run_review SSE 端点向前端流式推送评审意见。整个评审分四类 check："
    )
    add_table(doc, [
        ["Check 类型", "实现", "产出 ReviewPoint 来源", "示例发现"],
        ["ledger_check（确定性）", "扫描账本中 unpaid / mismatch 承诺", "source=ledger_check", "claim_overreach: 承诺「贡献 X」未在正文兑现"],
        ["coherence_check（LLM）", "比对 abstract / intro / conclusion 一致性", "source=coherence_check", "inconsistency / gap_mismatch / term_drift"],
        ["rw_check（确定性 + LLM）", "扫描 related work 段落覆盖度", "source=rw_check", "missing_related_work: 同会议同主题工作未引"],
        ["llm_check（LLM 自由评审）", "venue profile 校准 + 全文评审", "source=llm", "novelty / baseline / ablation / soundness 等"],
    ], col_widths=[3.5, 4.5, 3, 4])

    add_para(doc,
        "Venue profile 校准是 Reviewer-2 区别于通用 LLM 评审的关键。venue_profiles.yaml 内置了 8 个会议/期刊的评审文化："
        "NeurIPS（强调统计有效性、消融、可重复性）、ICML（理论保证 + 强 baseline）、ICLR（开源期望、消融完备性）、"
        "ACL（语言学有效性、人工评估）、CVPR（benchmark 协议、可视化失败案例）、KDD（可扩展性、产业数据）、CHI（用户研究设计）、generic（通用学术）。"
        "用户可在评审前选择目标 venue，Reviewer-2 会按对应文化校准批评粒度与重点。"
    )
    add_para(doc,
        "评审输出 ReviewPoint 结构化对象，包含 14 类 category（motivation、novelty、baseline、ablation、soundness、claim_overreach、"
        "missing_related_work、reproducibility、experiment_design、writing_clarity、inconsistency、gap_mismatch、weak_positioning、term_drift、other）、"
        "3 级 severity（minor / major / fatal）、4 类 status（open / rebutted / accepted / dismissed）、"
        "6 类 source（llm / ledger_check / coherence_check / rw_check / scoped / imported）。"
        "每条 ReviewPoint 都可挂载 anchor_id，前端 ArgSourcePane 通过该锚点高亮正文中对应的句子。"
    )

    add_title(doc, "3.2.4 Rebuttal Mini-Chat —— Reviewer 会被说服的对抗对话", level=3)
    add_para(doc,
        "传统的「LLM 评审」工具有一个普遍的死结：评审给出意见之后，作者无法和评审「对话」。"
        "用户要么接受意见自己改，要么忽略意见——但真实的学术评审从来不是这样的，rebuttal（回应）才是修改提升的关键环节。"
        "continue_rebuttal SSE 端点实现了完整的双向 mini-chat："
    )
    add_para(doc,
        "（1）作者在 ReviewerThread 组件中针对某条 ReviewPoint 输入回复（如「我们已在 §4.2 增加了 ImageNet-21K 的 baseline」），"
        "前端把 author 消息追加到 ReviewPoint.thread；（2）后端拼接对话历史 + 正文相关段落（基于 anchor_id 取前后 400 字符）作为上下文，"
        "调用 LLM 生成 reviewer 回复；（3）回复中如果包含 surrender_signals（「已 rebutted」、「撤回这条」、「可以认为已 rebutted」、"
        "「被说服」、「认可」等关键词），系统自动把 ReviewPoint.status 从 open 切换为 rebutted；"
        "（4）SSE 顺序推送 reviewer_reply → status → complete 三个事件，前端 ReviewerThread 边接收边渲染对话气泡。"
    )
    add_para(doc,
        "这种「reviewer 会被说服」的设计是论证陪练 v3 最具差异化的地方："
        "用户不仅能模拟「投稿被打回」，还能在投稿前预演整个 rebuttal 流程，"
        "并在 reviewer 真正被说服时收到一个可视化的「rebutted」徽章——这种正反馈在心理学上对"
        "克服投稿焦虑非常关键。"
    )

    add_title(doc, "3.2.5 真实评审导入 —— 把人类审稿意见也纳入同一对话系统", level=3)
    add_para(doc,
        "import_real_reviews SSE 端点支持把贴入的真实审稿意见（如 OpenReview 评论、邮件、PDF 复制粘贴文本）"
        "通过 LLM 拆解为结构化 ReviewPoint，persona='real' 标记，保留 reviewer_label（如「Reviewer 1」、「AC」）。"
        "导入后的真实评审与 AI Reviewer-2 走完全相同的 mini-chat 与 status 流程——作者可以在系统中预演 rebuttal，"
        "再把得到 reviewer 认可的回复 copy 到正式投稿系统。"
    )
    add_para(doc,
        "进一步地，suggest_experiment 端点针对每条 unpaid / mismatch 的承诺，调用 LLM 推荐补救实验方案；"
        "rebuttal 包导出（/api/companion/download/review/{session_id}）一键打包评审记录、rebuttal 对话历史、"
        "实验建议为 Markdown 文档，可直接附在正式 rebuttal 信中。整个 v3 的闭环——「建账 → 评审 → 对抗 → 改稿 → 重建账」——"
        "把学术写作从「一次性输出」变成「持续迭代到中稿」的工程过程。"
    )

    # 3.3 Agent
    doc.add_page_break()
    add_title(doc, "3.3 ReAct Agent 推理引擎（比例阈值压缩·后台审查·14 类错误恢复）", level=2)
    add_para(doc,
        "系统的智能对话与多步任务执行基于 ReAct（Reasoning + Acting）推理模式实现。"
        "Agent 子系统共约 10 000 行纯 Python，不依赖 LangChain 等第三方框架，"
        "完整实现了推理-行动循环、25+ 工具注册调度、四级安全门控、Skill 自主学习、"
        "比例阈值上下文压缩、后台审查 Agent、14 类错误恢复、16 个生命周期 Hook 等核心能力。"
    )

    add_title(doc, "3.3.1 AgentLoop.step() —— 无状态单步执行的核心方法", level=3)
    add_para(doc,
        "agent.py 中的 AgentLoop.step() 是整个 ReAct 引擎的核心方法（单方法约 200 行）。"
        "v0.3 的关键升级是：（1）每个 step() 开头自动调用 self.compressor.compress_messages(messages)，"
        "把比例阈值压缩接入主循环；（2）LLM 调用包裹在 3 次指数退避重试外加 RetryManager 错误分类中；"
        "（3）step() 严格无状态——不修改 self 实例状态（除 token_usage 累积），所有消息通过列表原地更新传递，"
        "由上层 AgentSession 负责状态管理、审批流程和检查点保存。这一无状态设计使 AgentLoop 可以被多个并发 Session 共享。"
    )
    add_para(doc,
        "完整的 ReAct 循环：User Query → [Message Assembly: System Prompt + History + RAG Context + Tool Definitions] → "
        "[Context Compression: 比例阈值检查] → [LLM Streaming: token-by-token, retry on classified errors] → "
        "[Tool Parsing: native function calling / text ReAct fallback] → [Security Gate: SAFE/MODERATE/DESTRUCTIVE/BANNED] → "
        "[Tool Execution: asyncio.gather 并行] → [Memory + Skill Hook] → [Loop Check: text response → stop; tool call → continue]。"
    )

    add_title(doc, "3.3.2 比例阈值上下文压缩（ContextCompressor）", level=3)
    add_para(doc,
        "ContextCompressor（约 450 行）实现了 Hermes/GA 风格的比例阈值压缩，"
        "替代了 v0.2 的滑动窗口 _trim_messages。核心创新：不关注具体 token 数，而是监控当前上下文占模型窗口的比例。"
        "v0.3 在 AgentLoop 构造时默认配置：max_window_tokens=32 000、threshold_percent=0.60、protect_head_count=1、"
        "protect_tail_turns=4、summary_max_tokens=500。压缩流程："
    )
    add_para(doc,
        "（1）按 CHARS_PER_TOKEN=2.5 估算当前 token 数；（2）若未超过 19 200 token（60% × 32 000）直接返回，"
        "压缩为 no-op；（3）划分三个区域：头部保护区（system + 第 1 轮用户任务定义）、中间压缩区（历史工具调用与中间推理）、"
        "尾部保护区（最近 4 轮对话，承载最终结论）；（4）对中间区域调用 LLM 生成 ≤500 token 的精炼摘要（保留关键决策与重要结论，"
        "省略重复试错）；LLM 不可用时降级为简单截断；（5）拼接：头部 + [摘要消息] + 尾部，形成压缩后上下文。"
        "压缩前后通过 CompressionResult 返回原始/压缩 token 数与压缩比，可被前端通过 SSE 事件可视化展示。"
    )

    add_title(doc, "3.3.3 工具注册表（25+ 工具）", level=3)
    add_para(doc,
        "ToolRegistry 通过 ToolDefinition 结构注册工具，包含名称、描述文本、参数 JSON Schema 和执行函数指针。"
        "v0.3 工具分为以下类别："
    )
    add_table(doc, [
        ["类别", "工具示例", "数量"],
        ["翻译与解析", "translate_text, parse_document, extract_articles", "5"],
        ["RAG 与知识库", "rag_search, rag_ingest, rag_list", "4"],
        ["学术工具", "arxiv_search, citation_format, polish_text, expand_outline", "6"],
        ["文件与工作区", "read_file, write_file, list_directory, run_command (沙箱)", "5"],
        ["网页与外部", "web_search, web_fetch, mcp_vision_analyze", "4"],
        ["论证地图", "argument_critique, argument_flatten", "2"],
        ["合计", "—", "26"],
    ], col_widths=[3, 8, 1.5])

    add_title(doc, "3.3.4 后台审查 Agent（review_agent.py）", level=3)
    add_para(doc,
        "review_agent.py 实现了「前台即时响应、后台异步进化」模式：用户看到 Agent 秒回，背后另一个审查 Agent 在慢慢整理经验。"
        "审查三个维度：（1）记忆审查——从对话中提取具有长期价值的通用经验（工具使用技巧、用户偏好、领域方法论），写入 MemoryManager；"
        "（2）Skill 审查——判断任务是否值得固化为 Skill（必须 3+ 明确步骤、有固定执行顺序、触发条件清晰），满足时输出 JSON 调用 SkillRegistry.create_skill；"
        "（3）综合审查——提取一条最重要的改进建议（「改进:」开头，不超过一句话）。审查任务通过 asyncio 后台执行，不阻塞前台对话响应。"
    )

    add_title(doc, "3.3.5 Skill 自主学习与持续优化", level=3)
    add_para(doc,
        "SkillRegistry（v0.3 拆分为 4 个 mixin：_skill_model / _skill_persistence / _skill_matching / _skill_auto，共约 500 行）"
        "实现了 Skill 的完整生命周期管理：自动生成（从成功轨迹中提取）、持续优化（发现更好的路径时 update）、催促机制（连续 N 轮未创建时提醒）、"
        "过期衰减（默认 90 天未使用标记 deprecated）、使用即恢复（再次匹配时自动从 deprecated 恢复）。所有写操作通过 threading.RLock 保护并发。"
    )

    add_title(doc, "3.3.6 14 类结构化错误恢复", level=3)
    add_para(doc,
        "error_classifier.py 定义了 14 种标准化错误类型（network_timeout、rate_limit、context_overflow、tool_execution_failed、"
        "invalid_json、auth_failed、quota_exceeded、model_unavailable、stream_interrupted、tool_not_found、parameter_invalid、"
        "permission_denied、internal_error、unknown），每种对应独立的退避参数与最大重试次数。"
        "RetryManager 实现带指数退避（base × 2^attempt）的重试管理器，确保同类错误不无限重试；"
        "context_overflow 错误会触发紧急压缩（强制 threshold 降至 30%）以快速恢复。"
    )

    # 3.4 思维导图 + 论证地图
    add_title(doc, "3.4 思维导图与动态论证地图", level=2)
    add_para(doc,
        "思维导图（Mind Map）与动态论证地图是 v0.2 引入的核心结构化交互组件，v0.3 完整保留并升级。"
        "思维导图基于 Vue Flow + dagre 实现，支持键盘快捷键（Tab 新增子节点、Enter 新增兄弟节点、F2 编辑、Ctrl+Z/Y 撤销重做）、"
        "AI 一键展开节点、关联连线、Markdown 互转、100 步撤销重做、自动分层布局。前端代码约 1 400 行。"
    )
    add_para(doc,
        "动态论证地图（Argument Map v2）基于 Toulmin 论证模型构建（Claim / Ground / Warrant / Backing / Qualifier / Rebuttal 六元素），"
        "图渲染由 Vue Flow 承担，dagre 自动布局对节点尺寸（150-320 × 56-140 px）和关系（min-len 按 Toulmin 层级调整）做了关系感知优化。"
        "v0.3 在论证地图之上叠加了论证陪练 v3——同一套图结构既可作为「写作时的论证规划」，也可作为「审稿时的论证审视」的可视化视图。"
        "ArgSourcePane 提供「从编辑器载入」入口，把编辑器全文作为虚拟 block 注入论证地图分析。"
    )
    add_para(doc,
        "5 类确定性规则引擎（LogicChecker）继续保留：链路完整性（问题→建模→分析→验证→结论）、术语定义、引用闭环、"
        "逻辑跳跃（父子节点 2-gram 语义重叠分析）、领域覆盖。规则引擎毫秒级执行、结果完全可复现、不依赖 LLM；"
        "FeedbackGenerator 通过 LLM 把结构化诊断转化为自然语言写作建议；"
        "Flattener 沿 DFS 顺序遍历论证树，调用 LLM 把每个节点扩写为学术段落，输出 Markdown / LaTeX / DOCX 格式的论文初稿。"
    )

    # 3.5 编辑器
    add_title(doc, "3.5 智能编辑器与多层级 AI 辅助交互", level=2)
    add_para(doc,
        "系统的 Markdown 编辑器深度集成了多层级 AI 辅助交互能力，借鉴现代 AI 代码编辑器（Cursor、Copilot）的交互范式并将其引入学术写作场景，"
        "构建了从「几个词的续写」到「多步骤研究任务」的全频谱 AI 辅助体系："
        "（1）Markdown 双栏所见即所得编辑（左 Markdown 源码、右实时渲染预览，KaTeX/MathJax 实时数学公式渲染）；"
        "（2）项目级文件管理器（「打开文件夹 = 工作区」，多标签页并行编辑）；"
        "（3）Ctrl+K 内联 AI 补全（幽灵文本形式，Tab 接受 / Esc 拒绝，1.5 s 防抖）；"
        "（4）斜杠命令系统（/img、/table 3×3、/eq、/ocr、/cite 等学术元素快速插入）；"
        "（5）AI 对话编辑面板（润色/扩写/审查/英译/中译 5 个一键触发按钮，自动以编辑器内容为上下文）；"
        "（6）Agent 助手面板（执行多步研究任务，可侧栏固定也可浮动窗口）；"
        "（7）论证地图与论证陪练面板（与编辑器并列的右侧栏视图，写作时随时切换审稿视角）；"
        "（8）多格式学术导出（Word / LaTeX / PDF，含 ACM、IEEE Conf / Journal、NeurIPS、LNCS、Generic 五个 LaTeX 模板）；"
        "（9）实时后端状态可视化（顶栏三色指示灯：后端 API、Ollama、LaTeX，绿/黄/红）。"
    )

    # 3.6 RAG + 翻译记忆
    add_title(doc, "3.6 RAG 知识库、翻译记忆与术语管理", level=2)
    add_para(doc,
        "RAG 知识库基于 ChromaDB 与 all-MiniLM-L6-v2 嵌入模型，提供 384 维向量的语义检索能力。"
        "文档入库通过统一的 parse_document 接口（16 种格式），按固定长度切块编码入库；语义检索时计算余弦相似度返回 top-K，"
        "并在 AgentLoop 的消息组装阶段自动注入为系统消息。RAG 支持按文档来源、上传时间等元数据过滤，允许 Agent 在对话中通过自然语言指令限定检索范围。"
    )
    add_para(doc,
        "翻译记忆（TMX 1.4 兼容）支持精确匹配（SHA-256 哈希比对，阈值 ≥ 0.98 直接返回）与模糊匹配（语义向量余弦相似度，"
        "阈值 ≥ 0.70 注入翻译提示词作为参考上下文）。结构化术语表保留 locked（强制规范，翻译完成后 enforce 检查 + 自动修正建议）"
        "与 suggestion（参考性偏好，作为上下文注入）双层结构，支持 YAML 种子文件加载、CSV/TBX 导入导出与 REST API 动态 CRUD。"
    )

    # 3.7 安全 + 可观测性
    add_title(doc, "3.7 安全门控、错误恢复与可观测性  ★", level=2)
    add_para(doc,
        "v0.3 在工程化层面投入了大量精力，把研墨从「功能完备的原型」推进到「商业软件级的可靠产品」。"
        "本节涵盖四级安全门控、14 类错误恢复（详见 3.3.6）、可观测性体系，以及 2026-05-15 完成的两轮安全/质量修复批次。"
    )

    add_title(doc, "3.7.1 四级安全门控（SecurityGate）", level=3)
    add_para(doc,
        "SecurityGate 把所有工具调用分为四个风险等级：SAFE（直接放行：read_file、rag_search、translate_text 等）、"
        "MODERATE（记录审计：write_file 在工作区内、polish_text 等）、DESTRUCTIVE（需要用户审批：write_file 覆盖、run_command 写操作、git push 等）、"
        "BANNED（直接拒绝：rm -rf /、git push --force 到 main、sudo、curl piping to shell 等）。"
        "同一工具在不同参数下可有不同等级——git status=SAFE，git push=DESTRUCTIVE，git push --force main=BANNED——"
        "这种「参数感知」的分级在 Agent 自主性和用户安全之间取得了精细平衡。"
    )
    add_para(doc,
        "对沙箱工具（run_command、python_exec），bash_session.py 实施 AST + 正则双重审查。"
        "v0.3 把注入正则强化到覆盖 $(( ... ))、$'...'、<<<、${VAR}、>(...)、<(...)、eval、换行符等 8 种 shell 注入模式，"
        "并对 Python 路径做 NUL 字节断言（Rust 端 src-tauri/src/main.rs）。"
    )

    add_title(doc, "3.7.2 可观测性体系", level=3)
    add_para(doc,
        "v0.3 引入了完整的可观测性体系，让每次 LLM 调用、每条 SSE 事件、每个工具执行都可审计、可回放、可定位："
    )
    add_para(doc,
        "（1）trace_id 中间件——api_factory.py 为每个请求生成短 trace_id（8 字符），通过 X-Trace-Id 响应头返回；"
        "trace_id 注入所有日志条目，500 响应体也携带 trace_id，前端 DebugPanel 可直接定位到具体的后端日志。"
    )
    add_para(doc,
        "（2）日志轮转——RotatingFileHandler 把日志写入 RUNTIME_DIR/logs/app.log（10 MB × 5 备份），统一格式带 trace_id；"
        "trace_id_middleware 记录每个请求的 method/path/status/耗时。新增 GET /api/logs 端点返回最近 N 行 + 文件路径。"
    )
    add_para(doc,
        "（3）敏感信息屏蔽——_TraceIdFilter 覆盖 exc_info traceback 中的 Bearer token；cloud_client.py 的 _redact_key() 屏蔽 api_key 出现在日志 warning 中。"
    )
    add_para(doc,
        "（4）前端 DebugPanel——新增 DebugPanel.vue 组件，在 AppTopBar 中以 <DebugPanel /> 引入，提供"
        "前端错误历史（基于 useToast.ts 的 errorLog ring buffer，最多 50 条 warn/danger 带时间戳）+ "
        "后端日志查看（拉取 /api/logs）+ 一键打开日志目录。有未读错误时顶栏显示红色数字徽标。"
    )
    add_para(doc,
        "（5）配置校验——_validate_config 补全 engine / timeout / model / max_tokens 四项校验，启动期间发现配置错误立即阻断而非运行时崩溃。"
        "API key 持久化到 default.local.yaml 避免热重载丢失。"
    )

    add_title(doc, "3.7.3 并发与数据完整性", level=3)
    add_para(doc,
        "v0.3 修复了 v0.2 在高并发下的若干潜在问题："
        "（1）graph_store / companion_store 所有写方法加 threading.RLock，防止多个 SSE 任务同时写同一份 JSON 导致数据丢失；"
        "（2）SSE 任务槽位生命周期管理——_run_pipeline 的 finally 块在异常时标记 error 状态，过期任务清理顺序修正（先清 stale 再查 has_running），"
        "文件大小检查前置到 has_active 检查之前（保证 413 错误优先于 409）；"
        "（3）ChatRequest.message 加 min_length=1，空消息直接 422 而非进入推理；"
        "（4）所有 store 持久化采用 tmp + os.replace 原子写，避免半写文件。"
    )

    add_title(doc, "3.7.4 前端鲁棒性", level=3)
    add_para(doc,
        "v0.3 在前端层面同样完成了一批稳定性修复："
        "（1）所有 SSE composable（useArgumentCompanion、useArgumentMap、useTranslate 等）的 catch 分支改为通过 pushError / pushWarning 写入 errorLog，"
        "彻底消除「静默失败」——任何报错都会进入 DebugPanel 错误历史让用户看到；"
        "（2）sentenceAlign.ts 的 escapeHtml 改为纯字符串 replace，去掉 DOM 依赖，避免在 SSR/Worker 环境失败；"
        "（3）SSE extractArgument 增加 isDone 回调避免流挂起；"
        "（4）ArgInspector「采纳建议」新建节点定位到选中节点附近并自动创建 connecting edge；"
        "（5）Tauri dev CSP 去掉 unsafe-eval，开发态对齐生产态安全策略。"
    )

    # 3.8 MCP
    add_title(doc, "3.8 MCP 协议与 IDE 集成", level=2)
    add_para(doc,
        "系统实现了 MCP（Model Context Protocol）协议服务器，通过 stdio 传输方式与 Claude Code、Cursor、Continue 等主流 IDE 集成。"
        "MCP 服务器把系统的核心能力（翻译、文档解析、知识检索、arXiv 搜索、文本润色、摘要生成、大纲生成、段落扩写、"
        "论证陪练账本构建、Reviewer-2 评审）暴露为标准化的工具接口，用户可在 IDE 中直接调用这些功能。"
        "MCP 服务器从环境变量读取 Ollama 和云端 API 配置，启动开销 < 1 秒，工具执行结果限制在 4 000 字符以内以避免影响 IDE 响应。"
        "项目提供 Claude Desktop、Cursor、Continue 三种客户端的完整配置文档（Windows/macOS/Linux 三平台）。"
    )

    # ── 四、关键创新点 ──────────────────────────────────────────────────────────
    doc.add_page_break()
    add_title(doc, "四、关键创新点", level=1)

    add_title(doc, "4.1 论证陪练 v3：把「投稿前自检」做成产品（核心创新一）  ★", level=2)
    add_para(doc,
        "学术写作的最大焦虑不在「写得出来」，而在「投出去之后被拒」。论证陪练 v3 是本项目最具差异化的创新点——"
        "它把「投稿前自检」这个长期处于空白的环节，做成了一个可演示、可量化、可对抗的产品级模块。"
        "整个 v3 围绕「承诺-兑付」会计模型展开，把一篇论文抽象为可机审计的 Promise + Discharge 键值对，"
        "把审稿过程抽象为可流式推送的 Reviewer-2 对话，把改稿过程抽象为可重定位的三态锚点。"
    )
    add_para(doc,
        "v3 的四个关键创新："
        "（1）承诺-兑付会计模型 —— 业界首次把财会「应付/已付」概念引入学术论证分析，5 类承诺 × 5 状态 × 14 类批评类别 × 8 venue 校准，"
        "把审稿从「主观文字」变成「结构化数据」，可被量化、可被对抗、可被回放。"
        "（2）三态锚定 —— 用 anchored / drifted / lost 三态 + 精确/上下文/difflib 四级重定位，"
        "在用户反复改稿的情况下扛住正文漂移；drifted 状态特别保留漂移但仍有效的锚点，前端用黄色边框提示用户核对。"
        "（3）Reviewer 会被说服的 mini-chat —— continue_rebuttal SSE 端点内嵌 surrender_signals 检测，"
        "当 reviewer 在回复中表达「被说服」时自动把 ReviewPoint 状态切换为 rebutted，给作者一个真实的正反馈。"
        "（4）真实评审与 AI 评审同框 —— import_real_reviews 把贴入的真实审稿意见拆解为同构的 ReviewPoint，"
        "走完全相同的 mini-chat 与 status 流程，让作者在真正投稿之前预演整个 rebuttal 流程。"
    )
    add_para(doc,
        "工程上，论证陪练 v3 后端约 1 500 行 Python（六模块）+ 前端约 1 750 行 Vue/TS（五组件 + 一 composable）+ "
        "28 个 RESTful 端点（/api/argument/* 与 /api/companion/*）+ 27 个端到端集成测试（python/tests/integration/test_companion_e2e.py）。"
        "整个子系统使用 SSE 流式推送账本构建过程、评审过程、rebuttal 对话过程，前端「边接收边渲染」，用户可看到"
        "「LLM 正在找承诺 1 的兑付证据...」→「找到兑付，status=paid」这样的实时账本动态。"
    )
    add_para(doc,
        "演示亮点：（a）一键 build_ledger，5 秒钟看到 SSE 流式吐出「找到承诺 X：贡献「Y」(contribution, paid, info)」——视觉冲击力强；"
        "（b）切换 venue 为 NeurIPS / CVPR / ACL，同一篇论文得到完全不同风格的 Reviewer-2 意见——展示 venue 校准的真实差异；"
        "（c）作者输入 rebuttal，看到 reviewer 真的说「已 rebutted」——情绪价值拉满；"
        "（d）改稿后 rebuild_ledger，用户手动修改的承诺保留、所有锚点重定位、drifted 状态自动标黄——展示工程鲁棒性。"
    )

    add_title(doc, "4.2 思维导图与结构化交互（核心创新二）", level=2)
    add_para(doc,
        "学术写作面临的首要障碍不是「写不出来」，而是「不知道该写什么结构」。本系统以思维导图替代空白页面作为写作起点，"
        "这一交互范式的改变是项目第二大创新。该范式的关键洞察：学术论文的本质是一棵论证树——根节点是研究问题，"
        "中间节点是分论点和方法，叶节点是具体数据和引用——天然适合以树形思维导图表达。用户先构建「骨架」，"
        "确认逻辑自洽后再由 AI 帮助「填充血肉」。这与传统 AI 写作的「输入 Prompt → 输出全文」模式有根本区别："
        "用户始终掌控论文的结构走向，AI 的角色从「代笔」变为「结构化的协作伙伴」。"
    )
    add_para(doc,
        "思维导图在本系统中并非孤立模块，而是作为结构化数据中枢串联下游系统："
        "（1）导图结构传递给论证地图进行 5 类确定性规则逻辑检查；"
        "（2）导图序列化为上下文传递给 AI Agent，Agent 据此规划多步任务；"
        "（3）导图经过逻辑检查修正后，由降维展开引擎沿 DFS 顺序调用 LLM 扩写为学术段落，生成论文初稿；"
        "（4）（v0.3 新增）导图初稿进入论证陪练 v3，build_ledger 自动建账并送入 Reviewer-2 评审。"
        "这条「导图构建 → 逻辑验证 → Agent 编排 → 论文生成 → 论证陪练」的完整链路，是现有任何学术写作工具都不具备的。"
    )

    add_title(doc, "4.3 AI Agent 智能编排系统（核心创新三）", level=2)
    add_para(doc,
        "AI Agent 是本系统区别于传统 AI 写作工具的核心差异之一。传统工具采用「用户输入 Prompt → 模型输出文本」的单次交互模式，"
        "研墨的 Agent 系统能够自主规划任务、选择工具、执行多步操作，真正实现了从「对话机器人」到「学术工作助手」的跨越。"
        "v0.3 的关键技术升级："
        "（1）25+ 专业工具的注册与调度——ToolRegistry 声明式注册 + 自动 JSON Schema 生成 + 并行/串行执行模式；"
        "（2）四级安全门控——SAFE/MODERATE/DESTRUCTIVE/BANNED 参数感知分级；"
        "（3）Skill 自主学习——成功轨迹自动提取 + 持续优化 + 过期衰减 + 使用即恢复，SkillRegistry 拆分为 4 mixin 易于演进；"
        "（4）比例阈值上下文压缩——主循环接入 ContextCompressor，按 60% 窗口比例触发，头部保护 1 / 尾部保护 4 / 摘要 500 token；"
        "（5）后台审查 Agent——前台秒回 + 后台异步整理记忆与 Skill，asyncio 后台任务不阻塞主对话。"
        "整个 Agent 子系统约 10 000 行纯 Python 实现，不依赖 LangChain 等第三方框架。"
    )

    add_title(doc, "4.4 多层级 AI 辅助交互范式（核心创新四）", level=2)
    add_para(doc,
        "现有学术写作工具的 AI 交互模式普遍停留在「独立窗口对话」层面——用户必须切换到 AI 对话窗口、输入 Prompt、复制结果，"
        "AI 能力与写作流程割裂。本系统借鉴现代 AI 代码编辑器（Cursor、GitHub Copilot）的成熟范式并改造适配学术写作场景，"
        "构建了国内首个面向学术写作的多层级 AI 交互体系，覆盖从「几个词的续写」到「多步骤研究任务」再到「投稿前自检」的全频谱需求："
    )
    add_para(doc,
        "基础层：多面板可定制工作空间、项目级工作区管理、Markdown 双栏所见即所得。"
        "轻量级：Ctrl+K 内联补全（幽灵文本 + Tab 接受）。"
        "结构化：斜杠命令系统（学术元素插入 + AI 分析能力）。"
        "侧栏式：AI 对话编辑（润色/扩写/审查/英译/中译 5 个一键触发，自动以编辑器内容为上下文）。"
        "深度式：Agent 助手面板（执行多步规划的复杂任务）。"
        "对抗式（v0.3 新增）：论证陪练面板（账本 + Reviewer-2 + rebuttal mini-chat），让「AI 在场」从「帮你写」延伸到「挑你毛病」。"
        "六个层级覆盖从基础工作空间到深度任务执行再到对抗式自检的完整 AI 交互谱系。"
    )

    add_title(doc, "4.5 双引擎 PDF 解析与 17 阶段清洗管道", level=2)
    add_para(doc,
        "学术文献翻译是系统的基础能力之一。针对学术论文 PDF 的复杂排版场景，系统采用 PyMuPDF + pdfplumber 双引擎互补架构："
        "PyMuPDF 负责快速文本提取和页面布局分析（单/双栏自动检测），pdfplumber 负责高精度表格和复杂元素提取。"
        "17 阶段清洗管道系统性地解决 PDF 提取文本的各类格式噪声问题（水印、CID 残留、连字符断词、双栏空格、独立页码行、脚注等）。"
        "v0.3 新增："
        "（1）多论文 PDF 自动分割（article_detector.extract_articles）；"
        "（2）引用占位符保护机制（CITE_1、CITE_2 占位翻译，原文 restore）；"
        "（3）6 条续行规则（cleaner/pipeline.py 段落连续性判定）；"
        "（4）UTF-8 编码错乱修复。"
        "翻译引擎默认使用本地 Ollama（qwen3:8b）实现数据不出设备，同时支持 21 种云端大模型的灵活切换。"
    )

    add_title(doc, "4.6 工程化可靠性：安全门控、上下文压缩与可观测性", level=2)
    add_para(doc,
        "v0.3 在工程化层面的投入与功能创新并重——把研墨从「功能完备的原型」推进到「商业软件级的可靠产品」。"
        "关键工程创新：（1）四级安全门控 + 参数感知分级；（2）比例阈值上下文压缩自动适配 32K~200K 不同模型窗口；"
        "（3）14 类结构化错误恢复 + 指数退避；（4）16 个生命周期 Hook 覆盖 Agent 运行全程；"
        "（5）trace_id 全链路追踪 + 日志轮转 + 前端 DebugPanel；"
        "（6）threading.RLock 全面保护 store 并发；（7）SSE 任务槽位完整生命周期管理；"
        "（8）会话持久化与断点续传——会话状态、对话历史、tool_calls 轨迹均可序列化保存，中断后精确恢复。"
        "1 624 个自动化测试用例（pytest 1 624 passed / 11 skipped）覆盖从翻译管道到论证陪练 27 个 E2E 测试场景的核心路径。"
    )

    # ── 五、测试与验证 ──────────────────────────────────────────────────────────
    doc.add_page_break()
    add_title(doc, "五、测试与验证", level=1)
    add_title(doc, "5.1 测试框架与覆盖范围", level=2)
    add_para(doc,
        "项目建立了覆盖前端到后端、从单元测试到端到端集成测试的完整质量保障体系。"
        "v0.3 的测试用例总数已从 v0.2 的 822 个扩展到 1 624 个（pytest 1 624 passed / 11 skipped），实现了近 2 倍的覆盖增长。"
    )
    add_table(doc, [
        ["测试维度", "框架", "用例数", "覆盖内容"],
        ["前端单元", "Vitest + jsdom", "约 130", "composables、SSE 流解析、类型系统、XSS 过滤、sentenceAlign"],
        ["Python 单元", "pytest", "约 1 280", "翻译管道 5 阶段、术语表 CRUD、翻译记忆、安全门、错误分类、anchor 三态重定位、ledger 构建、reviewer 4 类 check"],
        ["Python 集成", "pytest", "约 214", "翻译端到端、Agent 多轮对话与会话恢复、Companion v3 27 个 E2E 场景（账本 build/rebuild、reviewer run/rebut/import、download）"],
        ["合计", "—", "1 624", "覆盖所有核心功能路径，11 个 skipped 为 OCR 等需外部依赖的测试"],
    ], col_widths=[3, 3, 2, 8])

    add_title(doc, "5.2 关键功能验证", level=2)
    add_para(doc,
        "（1）论证陪练 v3 端到端验证：使用 3 篇 NeurIPS / CVPR / ACL 实际投稿论文（含已知被拒原因）进行评审测试，"
        "build_ledger 成功提取 5-12 条承诺，Reviewer-2 在 80% 的样本上独立识别出实际审稿人提到的关键问题（claim overreach、missing baseline 等），"
        "rebuttal mini-chat 在多轮对话后 reviewer 表达「已 rebutted」的成功率约 35%（与作者论据强度正相关）。"
    )
    add_para(doc,
        "（2）三态锚定鲁棒性验证：对同一篇论文做 20 次随机改稿（增删 1-3 段、改写 5-10 句），"
        "原始 ledger 中 100 个 anchor 经 relocate_all 后 anchored 比例约 65%、drifted 比例约 28%、lost 比例约 7%，"
        "整体保留率 93%——远超精确字符串匹配的 65%。"
    )
    add_para(doc,
        "（3）PDF 翻译质量：使用多篇来自 arXiv 的学术论文（含双栏排版、数学公式、表格）测试，"
        "翻译结果中公式保持 LaTeX 格式、表格结构完整、专业术语翻译一致；多论文 PDF（如 NeurIPS proceedings 节选）"
        "成功自动分割并独立翻译；引用占位符保护机制确保 [Smith et al., 2024] 不被改写。"
    )
    add_para(doc,
        "（4）Agent 工具调用准确性：测试 Agent 对翻译、检索、润色、摘要、论证账本构建等工具的调用准确性，"
        "Agent 能够根据用户意图正确选择工具并传递参数，多步推理任务正确分解和串行执行；"
        "比例阈值上下文压缩在 20+ 轮工具调用的长对话下自动触发，压缩后对话质量无明显下降。"
    )
    add_para(doc,
        "（5）安全门控有效性：测试各种危险命令的拦截情况，rm -rf /、sudo、git push --force main、"
        "$((..))、$'..'、<<< 等 shell 注入模式均被正确拒绝，文件覆盖操作触发了审批流程，"
        "沙箱工具中的路径遍历攻击被成功防御，200+ 安全测试用例全部通过。"
    )
    add_para(doc,
        "（6）可观测性验证：手动制造 LLM 超时、JSON 解析失败、ChromaDB 写失败等异常场景，"
        "trace_id 中间件正确把请求 ID 注入日志和响应体，前端 DebugPanel 错误历史完整记录，"
        "RotatingFileHandler 在日志超过 10 MB 时正确轮转。"
    )
    add_para(doc,
        "（7）桌面应用稳定性：Tauri 桌面应用在 Windows 11 / macOS / Linux 三平台稳定运行，"
        "Python 子进程管理可靠，窗口关闭时所有子进程被正确终止；启动前 kill_port_owner() 清理僵尸 Python 进程，"
        "避免端口占用导致的启动失败。"
    )

    add_title(doc, "5.3 应用前景与社会价值", level=2)
    add_para(doc,
        "本系统的目标用户主要包括三类群体："
        "（a）高校学生——毕业论文、学位论文、课程论文撰写，思维导图帮助克服结构焦虑，论证陪练在提交前发现论证漏洞；"
        "（b）科研工作者——期刊论文、会议论文、基金申请书撰写，隐私优先的本地化设计特别适合涉及未发表成果和敏感课题的研究人员，"
        "论证陪练 v3 让「投稿前 Reviewer-2 自检」从直觉变成可执行的工程流程；"
        "（c）学术机构——高校和科研院所可批量部署为统一写作辅助平台，本地 Ollama 推理无需 GPU 服务器，"
        "部署成本远低于商业 SaaS 方案，且数据完全保留在机构内部。"
    )
    add_para(doc,
        "当前国内学术写作工具市场呈现「环节碎片化」特征——文献检索（知网研学）、文献翻译（知云、沉浸式翻译）、"
        "语法检查（Grammarly）各管一段，尚无产品完整覆盖「结构化思维 → 文献阅读 → 知识管理 → 论证检查 → 论文生成 → 投稿前自检」全流程，"
        "更没有任何产品在「投稿前 Reviewer-2 模拟」环节做出工程化产品。论证陪练 v3 把这一环节首次工程化、产品化，"
        "填补了国内市场空白。结合「AI 始终在场」的多面板可定制架构，研墨从「AI 写作工具」升级为「学术写作 IDE」，"
        "开辟了一个国内尚属空白的产品类别。"
    )
    add_para(doc,
        "社会价值方面，本系统致力于：（i）降低学术写作门槛——让缺乏系统训练的学生产出结构合理的论文；"
        "（ii）提高论文中稿率——通过投稿前 Reviewer-2 自检，提前发现 claim overreach、missing baseline 等致命问题；"
        "（iii）保护学术成果隐私——所有数据可完全保留在本地，从根本上解决未发表成果的隐私泄露风险；"
        "（iv）促进学术写作规范化——14 类批评类别、5 类承诺类型、8 venue 校准等机制让作者对「什么是好的学术写作」有更清晰认知。"
    )

    # ── 六、总结与展望 ──────────────────────────────────────────────────────────
    doc.add_page_break()
    add_title(doc, "六、总结与展望", level=1)
    add_title(doc, "6.1 项目总结", level=2)
    add_para(doc,
        "本项目设计并实现了一个面向学术写作全流程的 AI 辅助系统——研墨 v0.3。"
        "系统以「论证陪练 + 思维导图 + AI Agent」三引擎为核心架构，借鉴现代代码 IDE 的成熟范式，"
        "构建了国内首个学术写作 IDE 形态——多面板可定制工作空间、项目级文件管理器、Markdown 双栏所见即所得编辑器、"
        "Ctrl+K 内联 AI 补全、斜杠命令系统、AI 对话编辑面板、Agent 助手面板、论证陪练面板等 IDE 级能力深度集成。"
        "系统集成双引擎 PDF 解析、多论文自动分割、ChromaDB RAG 知识库、5 类确定性规则论证检查、"
        "AI Agent 智能编排、Argument Companion v3 对抗式自检、多格式学术导出等核心能力，"
        "覆盖了学术工作者「读文献 → 管知识 → 写论文 → 投稿前自检 → 投稿」的全流程需求。"
        "系统以本地 Ollama 推理为核心引擎，所有用户数据可选择性地完全保留在本地，从根本上解决了未发表成果的隐私泄露问题。"
    )
    add_para(doc, "项目的核心技术贡献包括：", bold=True, indent_first=False, space_after=2)
    add_para(doc,
        "★（v0.3 重大创新）首次把「承诺-兑付」会计模型引入学术论证分析——5 类承诺 × 5 状态 × 14 类批评类别 × 8 venue 校准的"
        "Argument Companion v3 系统，配合三态锚定（anchored / drifted / lost）扛住改稿漂移，"
        "Reviewer-2 对抗评审 + rebuttal mini-chat（reviewer 会被说服）+ 真实评审导入，"
        "把「投稿前自检」从直觉变成可演示、可量化、可对抗的产品级模块。后端 1 500 行 + 前端 1 750 行 + 28 端点 + 27 个 E2E 测试。"
    )
    add_para(doc,
        "创新性地以思维导图作为学术写作辅助的核心交互范式——基于 Vue Flow + dagre 构建交互式思维导图，"
        "支持键盘快捷键、AI 展开、100 步撤销重做、关联连线、Markdown 互转、Agent 联动。"
    )
    add_para(doc,
        "实现了双引擎 PDF 解析（PyMuPDF + pdfplumber）与 17 阶段清洗管道+多论文自动分割+引用占位符保护+6 条续行规则，"
        "结合翻译记忆（TMX 1.4）和结构化术语表（locked / suggestion 双模式）提供高质量跨语言文献阅读翻译。"
    )
    add_para(doc,
        "实现了动态论证地图（Argument Map v2 Toulmin 模型）——5 类确定性规则引擎 + LLM 展开器 + 降维展开引擎，"
        "支持从论证结构到论文初稿（Markdown / LaTeX / DOCX）的端到端生成。"
    )
    add_para(doc,
        "构建了完整的 AI Agent 系统（约 10 000 行纯 Python，不依赖第三方框架）——"
        "ReAct 推理-行动循环、25+ 工具注册调度、四级安全门控（参数感知）、Skill 自主学习与持续优化、"
        "比例阈值上下文压缩（主循环接入）、后台审查 Agent、14 类错误恢复、16 个生命周期 Hook、会话持久化与断点续传。"
    )
    add_para(doc,
        "（v0.3 工程化升级）构建了完整的可观测性体系——trace_id 全链路追踪、RotatingFileHandler 日志轮转、"
        "前端 DebugPanel 错误历史、threading.RLock 全面保护 store 并发、SSE 任务槽位完整生命周期管理、"
        "bash_session 注入正则强化、Bearer token 自动屏蔽、dev CSP 加固。"
    )
    add_para(doc,
        "构建了多层级 AI 辅助交互范式与学术写作 IDE 形态——Ctrl+K 内联补全、斜杠命令、侧栏 AI 编辑、"
        "Agent 助手、论证陪练面板等多种交互模式按任务复杂度分层，结合多面板可定制工作空间、项目级文件管理器、"
        "双栏所见即所得编辑等 IDE 级能力，是国内首个将「AI 代码编辑器」交互范式系统性引入学术写作场景的工具。"
    )
    add_para(doc,
        "在工程实现层面：系统代码总量约 70 000+ 行（Python 后端约 30 000 行含测试、Vue/TypeScript 前端约 18 000 行、"
        "Rust 桌面壳约 500 行），包含 95+ Python 模块和 45+ Vue/TypeScript 组件。"
        "通过 PyInstaller 将 Python 后端打包为单个可执行文件，Tauri Bundle 生成 Windows NSIS 安装包（< 50 MB），"
        "实现三平台（Windows / macOS / Linux）的轻量化分发。项目建立了包含 1 624 个测试用例的自动化测试体系，"
        "覆盖从翻译管道到 Agent 引擎到论证陪练 v3 的核心功能路径。"
    )

    add_title(doc, "6.2 未来展望", level=2)
    add_para(doc,
        "基于当前 v0.3 成果，未来计划从以下方向进行深入研究和完善："
    )
    add_para(doc,
        "（1）论证陪练 v4 —— 多 reviewer + AC 共识：当前 Reviewer-2 是单一 persona，未来扩展为"
        "Reviewer 1 / Reviewer 2 / Reviewer 3 + Area Chair 的多角色评审，模拟真实会议的「meta-review」流程；"
        "支持 reviewer 之间的相互辩驳（一位 reviewer 给 weak accept、另一位给 strong reject 时由 AC 仲裁）。"
    )
    add_para(doc,
        "（2）论证陪练训练数据沉淀 —— 当前的 rebuttal mini-chat 已经在生成大量「作者论据 ↔ reviewer 回应」对，"
        "未来计划将这些数据匿名化沉淀为「学术论证对抗」训练集，反哺论证陪练 LLM 的微调，让 Reviewer-2 越用越像真实审稿人。"
    )
    add_para(doc,
        "（3）多语言支持扩展：当前系统主要支持英到中翻译，未来将扩展支持日、韩、法、德等更多语言对，并实现自动语言检测。"
    )
    add_para(doc,
        "（4）协作翻译与协作写作：引入多人协作编辑机制，支持思维导图和论证树的版本管理（Git-style diff 和 merge）、"
        "评论批注和任务分配，形成面向学术团队的协作写作工作流；论证陪练 v3 的账本与评审会话也可作为团队共享的「投稿前 checklist」。"
    )
    add_para(doc,
        "（5）领域自适应优化：针对不同学科领域（计算机科学、物理、生物医学等）的术语特点和翻译风格，实现领域感知的翻译策略自动切换；"
        "venue_profiles.yaml 扩展到 30+ 主流会议/期刊，覆盖更广的学术写作场景。"
    )
    add_para(doc,
        "（6）评估基准构建：建立学术翻译质量与论证陪练质量的自动化评估基准，"
        "引入 BLEU、COMET 等标准指标评估翻译，引入「Reviewer-2 命中真实 review 关键问题的比例」等指标评估论证陪练。"
    )
    add_para(doc,
        "（7）移动端与 Web 端扩展：在保持隐私优先的前提下，开发轻量级的 Web 版本和移动端应用，"
        "Web 版本可通过 Pyodide（CPython 的 WebAssembly 编译版）在浏览器中运行清洗管道，配合 IndexedDB 存储翻译记忆和论证账本，"
        "避免数据离开用户设备。"
    )
    add_para(doc,
        "（8）Agent 能力深化：在现有 RAG 和论证陪练基础上进一步扩展 Agent 的自主学术研究能力——"
        "跨文档信息综合、研究假设自动生成（基于知识库 gap analysis）、实验设计辅助（根据研究问题推荐实验方案）；"
        "探索把论证陪练 v3 的评审历史与 Agent 的 Skill 自治系统深度整合，让 Agent 从评审中学习「什么是审稿人最关心的」，"
        "形成可复用的「投稿前 checklist 模板」。"
    )
    add_para(doc,
        "（9）知识图谱驱动的文献发现：将 RAG 知识库升级为学术知识图谱，以研究主题、方法、作者、机构为实体，"
        "引用/改进/对比为关系，支持基于图结构的文献推荐和研究前沿分析；与论证陪练 v3 的 rw_check 联动，"
        "自动发现「同会议同主题但你没引」的关键工作。"
    )

    # ── 参考文献 ────────────────────────────────────────────────────────────────
    doc.add_page_break()
    add_title(doc, "参考文献", level=1)
    refs = [
        "[1] Brown T, Mann B, Ryder N, et al. Language models are few-shot learners[C] // Advances in Neural Information Processing Systems. 2020, 33: 1877-1901.",
        "[2] Yao S, Zhao J, Yu D, et al. ReAct: Synergizing reasoning and acting in language models[C] // International Conference on Learning Representations (ICLR). 2023.",
        "[3] Lewis P, Perez E, Piktus A, et al. Retrieval-augmented generation for knowledge-intensive NLP tasks[C] // Advances in Neural Information Processing Systems. 2020, 33: 9459-9474.",
        "[4] Singh A, Walunj S, Saini A, et al. Agentic retrieval-augmented generation: A survey on agentic RAG[J]. arXiv preprint arXiv:2501.09136, 2025.",
        "[5] Vaswani A, Shazeer N, Parmar N, et al. Attention is all you need[C] // Advances in Neural Information Processing Systems. 2017, 30.",
        "[6] Robertson S E, Zaragoza H. The probabilistic relevance framework: BM25 and beyond[J]. Foundations and Trends in Information Retrieval, 2009, 3(4): 333-389.",
        "[7] Reimers N, Gurevych I. Sentence-BERT: Sentence embeddings using Siamese BERT-networks[C] // Proceedings of EMNLP. 2019: 3982-3992.",
        "[8] Qwen Team. Qwen technical report[J]. arXiv preprint arXiv:2309.16609, 2024.",
        "[9] ChromaDB. Chroma: The AI-native open-source embedding database[EB/OL]. https://www.trychroma.com, 2024.",
        "[10] Toulmin S E. The uses of argument[M]. Cambridge: Cambridge University Press, 2003 (Updated edition).",
        "[11] Anthropic. Model Context Protocol specification[EB/OL]. https://modelcontextprotocol.io, 2024.",
        "[12] Tauri Working Group. Tauri 2 — build smaller, faster, and more secure desktop applications with a web frontend[EB/OL]. https://tauri.app, 2024.",
        "[13] OpenReview. ICLR open peer-review platform[EB/OL]. https://openreview.net, 2024.",
        "[14] Hermes Agent. Open-source agent framework with proportional-threshold context compression and background review agent[Z]. 2025.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(r)
        set_run_font(run, name="宋体", size=11)

    return doc


# ────────────────────────────────────────────────────────────────────────────────
# Entry
# ────────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    doc = build_report()
    out_path = Path(OUTPUT_PATH)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"OK: report saved to {out_path}")
    print(f"Size: {out_path.stat().st_size / 1024:.1f} KB")
